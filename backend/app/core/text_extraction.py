"""Text extraction from resume files.

Supports extracting text from PDF, DOCX, and plain text files.
"""

import logging

logger = logging.getLogger(__name__)


# Supported MIME types
SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}


def is_supported_mime_type(mime_type: str) -> bool:
    """Check if a MIME type is supported for text extraction."""
    return mime_type in SUPPORTED_MIME_TYPES


def get_supported_mime_types() -> list[str]:
    """Get list of supported MIME types."""
    return list(SUPPORTED_MIME_TYPES.keys())


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file.
    
    Uses pypdf to extract text from PDF documents.
    
    Args:
        file_bytes: Raw PDF file content
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If extraction fails
    """
    try:
        from pypdf import PdfReader
        from io import BytesIO

        reader = PdfReader(BytesIO(file_bytes))
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from PDF ({len(reader.pages)} pages)")
        return full_text.strip()

    except ImportError:
        raise ValueError("pypdf is not installed. Install it with: pip install pypdf")
    except Exception as e:
        logger.exception("PDF text extraction failed")
        raise ValueError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file.
    
    Uses python-docx to extract text from Word documents.
    
    Args:
        file_bytes: Raw DOCX file content
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If extraction fails
    """
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(file_bytes))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        return full_text.strip()

    except ImportError:
        raise ValueError("python-docx is not installed. Install it with: pip install python-docx")
    except Exception as e:
        logger.exception("DOCX text extraction failed")
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file.
    
    Attempts to decode the file as UTF-8, falling back to other encodings.
    
    Args:
        file_bytes: Raw text file content
        
    Returns:
        Decoded text content
        
    Raises:
        ValueError: If decoding fails
    """
    # Try common encodings
    encodings = ["utf-8", "utf-16", "latin-1", "ascii"]
    
    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding)
            logger.info(f"Extracted {len(text)} characters from text file (encoding: {encoding})")
            return text.strip()
        except (UnicodeDecodeError, LookupError):
            continue
    
    raise ValueError("Failed to decode text file - unsupported encoding")


def extract_text_from_file(file_bytes: bytes, mime_type: str) -> str:
    """Extract text from a file based on its MIME type.
    
    Args:
        file_bytes: Raw file content
        mime_type: The MIME type of the file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If MIME type is not supported or extraction fails
    """
    if not is_supported_mime_type(mime_type):
        raise ValueError(
            f"Unsupported file type: {mime_type}. "
            f"Supported types: {', '.join(get_supported_mime_types())}"
        )

    file_type = SUPPORTED_MIME_TYPES[mime_type]
    
    if file_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        return extract_text_from_docx(file_bytes)
    elif file_type == "txt":
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


